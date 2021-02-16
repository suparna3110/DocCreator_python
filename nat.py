import emoji
import docx
import string 
import random 
import os 
import functools
import operator

# instance creation of a word document
document = docx.Document()
check = 0
# function for finding emoji in the string
# function for finding emoji in the string
def is_emoji(s,emojiEnabler):
    global check
    print(check)
    emojis = ["👍", "😁", "😊", "😛", "😔"]    
    demoji = emoji.demojize(s)
    if s == demoji:
        return False
    else:
        if emojiEnabler:
            if (s in emojis):
                return True
            else:
                print("Unknown Emoji found. Document will not contain unknown emoji.")
                check = 1
                return False
    return True

# function for writing string in the word document(with prdefined emoji)
def writing_word(text, emojiEnabler):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = docx.shared.Pt(11)
    paragraph = document.add_paragraph()
    for word in text:
        if is_emoji(word, emojiEnabler):
            print("ok")
            sentence = paragraph.add_run(word)
            sentence.font.name = 'Menlo'
            sentence.font.size = docx.shared.Pt(16)
        else:
            if check:
                continue
            paragraph.add_run(word)
        paragraph.add_run(" ")


def main_function():
    try:
        print("Select any one:" + "\n" + "1. Define own Emoji." + "\n" + "2. Use Pre Define Emoji." + "\n")
        selection = int(input("Enter Input: "))
        if selection == 1:
            text = input("Enter The Text: ")
            em_split_emoji = emoji.get_emoji_regexp().split(text)
            em_split_whitespace = [substr.split() for substr in em_split_emoji]
            text = functools.reduce(operator.concat, em_split_whitespace)
            # text = text.split()
            emojiEnabler = False
        elif selection == 2:
            text = input("Enter The Text: (You can use only these emoji's in your text- 👍, 😁, 😊, 😛, 😔)")  
            em_split_emoji = emoji.get_emoji_regexp().split(text)
            em_split_whitespace = [substr.split() for substr in em_split_emoji]
            text = functools.reduce(operator.concat, em_split_whitespace)  
            # text = text.split()

            emojiEnabler = True
        writing_word(text, emojiEnabler)
        res = ''.join(random.choices(string.ascii_uppercase + string.digits, k=1))
        dir_path = str(os.path.dirname(os.path.realpath(__file__)))
        print(dir_path)
        dir_path = dir_path + "/"
        document.save(dir_path + "Result" + res + ".docx")
    except:
        print("Problem in saving the document")

main_function()


